"""
文档解析服务
支持：PDF（优先 MinerU API，降级 PyPDF）、Word（python-docx）、URL（Jina Reader）
"""
import os
import httpx
import logging
from typing import Optional
from ..config import settings

logger = logging.getLogger(__name__)


async def parse_document(file_path: str, source_type: str) -> str:
    """
    统一入口：根据 source_type 分派解析器
    返回提取的纯文本内容（Markdown 格式）
    """
    if source_type == "pdf":
        return await parse_pdf(file_path)
    elif source_type == "word":
        return parse_word(file_path)
    elif source_type == "url":
        return await parse_url(file_path)
    else:
        raise ValueError(f"不支持的文档类型: {source_type}")


# ─────────────────────────────────────────────
#  PDF 解析
# ─────────────────────────────────────────────

async def parse_pdf(file_path: str) -> str:
    """优先调用 MinerU API，失败则降级到本地 PyPDF"""
    if settings.MINERU_API_KEY:
        try:
            return await _parse_pdf_mineru(file_path)
        except Exception as e:
            logger.warning(f"MinerU 解析失败，降级到 PyPDF: {e}")
    return _parse_pdf_local(file_path)


async def _parse_pdf_mineru(file_path: str) -> str:
    """调用 MinerU 云端 API 解析 PDF → Markdown"""
    async with httpx.AsyncClient(timeout=120) as client:
        with open(file_path, "rb") as f:
            files = {"file": (os.path.basename(file_path), f, "application/pdf")}
            headers = {"Authorization": f"Bearer {settings.MINERU_API_KEY}"}
            resp = await client.post(
                f"{settings.MINERU_API_URL}/file-urls/batch",
                files=files,
                headers=headers,
            )
            resp.raise_for_status()
            data = resp.json()
            # MinerU 返回 task_id，需要轮询结果
            task_id = data.get("data", {}).get("task_id")
            if not task_id:
                raise ValueError("MinerU 未返回 task_id")
            return await _poll_mineru_result(client, task_id, headers)


async def _poll_mineru_result(client: httpx.AsyncClient, task_id: str, headers: dict) -> str:
    import asyncio
    for _ in range(60):  # 最多等待 120 秒
        await asyncio.sleep(2)
        resp = await client.get(
            f"{settings.MINERU_API_URL}/extract-results/{task_id}",
            headers=headers,
        )
        resp.raise_for_status()
        data = resp.json()
        state = data.get("data", {}).get("state")
        if state == "done":
            # 拼接所有页的 markdown
            pages = data["data"].get("full_zip_url", "")
            # 简化处理：直接从 markdown 字段获取
            markdown_list = data["data"].get("detail", [])
            return "\n\n".join(
                item.get("md_content", "") for item in markdown_list if item.get("md_content")
            )
        elif state == "failed":
            raise ValueError(f"MinerU 任务失败: {data}")
    raise TimeoutError("MinerU 解析超时")


def _parse_pdf_local(file_path: str) -> str:
    """本地 PyPDF 解析（备用方案）"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        texts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                texts.append(f"## 第 {i + 1} 页\n\n{text.strip()}")
        return "\n\n".join(texts)
    except Exception as e:
        logger.error(f"PyPDF 解析失败: {e}")
        raise


# ─────────────────────────────────────────────
#  Word 解析
# ─────────────────────────────────────────────

def parse_word(file_path: str) -> str:
    """使用 python-docx 解析 .docx 文件"""
    from docx import Document
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 识别标题样式
        if para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "")
            try:
                level_num = int(level)
                parts.append(f"{'#' * level_num} {text}")
            except ValueError:
                parts.append(f"## {text}")
        else:
            parts.append(text)

    # 提取表格内容
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


# ─────────────────────────────────────────────
#  URL 解析（Jina Reader）
# ─────────────────────────────────────────────

async def parse_url(url: str) -> str:
    """
    使用 Jina Reader API (r.jina.ai) 将网页转换为 Markdown
    无需 API Key，直接访问
    """
    jina_url = f"https://r.jina.ai/{url}"
    async with httpx.AsyncClient(timeout=30) as client:
        headers = {
            "Accept": "text/plain",
            "X-Return-Format": "markdown",
        }
        resp = await client.get(jina_url, headers=headers, follow_redirects=True)
        resp.raise_for_status()
        return resp.text


# ─────────────────────────────────────────────
#  文本分块
# ─────────────────────────────────────────────

def split_text_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    将长文本按语义分块
    优先在段落边界切分，保留上下文重叠
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
        length_function=len,
    )
    chunks = splitter.split_text(text)
    # 过滤过短的无效块
    return [c.strip() for c in chunks if len(c.strip()) >= 50]
